"""
GEPPETTO 3: DOCUMENT LOADERS
=============================
Loader registry for the incremental KB sync.

Each loader takes an absolute file path and returns:
  {
    "text":     str,          # extracted plain text
    "metadata": {
        "source":        str,  # filename (relative, set by caller)
        "file_type":     str,  # extension without dot
        "modified_date": str,  # YYYY-MM-DD of file mtime
        "char_count":    int,
    }
  }
  Returns None if the file cannot be read (logs the reason).

Supported:
  .md / .txt  — built-in
  .pdf        — pymupdf  (pip install pymupdf)
  .docx       — python-docx  (pip install python-docx)
  .xlsx       — openpyxl  (optional, flag-gated — set ENABLE_XLSX=1)

Usage:
  from doc_loaders import load_document, SUPPORTED_EXTENSIONS
  result = load_document("/path/to/report.pdf")
  if result:
      print(result["text"][:200])
"""

import os
import logging
from datetime import date, datetime

logger = logging.getLogger(__name__)

# Set to True to enable xlsx loading (requires openpyxl)
ENABLE_XLSX = os.getenv("ENABLE_XLSX", "0") == "1"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _mtime_date(path: str) -> str:
    """Return the file's last-modified date as YYYY-MM-DD."""
    try:
        ts = os.path.getmtime(path)
        return datetime.fromtimestamp(ts).date().isoformat()
    except OSError:
        return date.today().isoformat()


def _base_meta(path: str, file_type: str, text: str) -> dict:
    return {
        "source":        os.path.basename(path),
        "file_type":     file_type,
        "modified_date": _mtime_date(path),
        "char_count":    len(text),
    }


# ---------------------------------------------------------------------------
# Individual loaders
# ---------------------------------------------------------------------------

def _load_text(path: str) -> dict | None:
    """Load .md or .txt files."""
    try:
        with open(path, encoding="utf-8", errors="replace") as f:
            text = f.read()
        ext = os.path.splitext(path)[1].lstrip(".") or "txt"
        return {"text": text, "metadata": _base_meta(path, ext, text)}
    except Exception as e:
        logger.warning("text loader failed for %s: %s", path, e)
        return None


def _load_pdf(path: str) -> dict | None:
    """Load .pdf using pymupdf (fitz). Extracts text from all pages."""
    try:
        import fitz  # pymupdf
    except ImportError:
        logger.error("pymupdf not installed. Run: pip install pymupdf")
        return None
    try:
        doc = fitz.open(path)
        pages = []
        for page in doc:
            pages.append(page.get_text())
        doc.close()
        text = "\n\n".join(pages).strip()
        if not text:
            logger.warning("pdf loader: no text extracted from %s (scanned?)", path)
            return None
        return {"text": text, "metadata": _base_meta(path, "pdf", text)}
    except Exception as e:
        logger.warning("pdf loader failed for %s: %s", path, e)
        return None


def _load_docx(path: str) -> dict | None:
    """Load .docx using python-docx. Extracts paragraphs and table cells."""
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return None
    try:
        doc = Document(path)
        parts = []
        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = "  |  ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)
        text = "\n\n".join(parts)
        if not text:
            logger.warning("docx loader: no text extracted from %s", path)
            return None
        return {"text": text, "metadata": _base_meta(path, "docx", text)}
    except Exception as e:
        logger.warning("docx loader failed for %s: %s", path, e)
        return None


def _load_xlsx(path: str) -> dict | None:
    """Load .xlsx using openpyxl. Converts sheets to text tables (flag-gated)."""
    if not ENABLE_XLSX:
        logger.debug("xlsx loading disabled. Set ENABLE_XLSX=1 to enable.")
        return None
    try:
        import openpyxl
    except ImportError:
        logger.error("openpyxl not installed. Run: pip install openpyxl")
        return None
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"## Sheet: {sheet_name}")
            for row in ws.iter_rows(values_only=True):
                row_text = "  |  ".join(
                    str(cell) for cell in row if cell is not None
                )
                if row_text.strip():
                    parts.append(row_text)
        wb.close()
        text = "\n".join(parts)
        if not text.strip():
            return None
        return {"text": text, "metadata": _base_meta(path, "xlsx", text)}
    except Exception as e:
        logger.warning("xlsx loader failed for %s: %s", path, e)
        return None


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------

_LOADERS = {
    ".md":   _load_text,
    ".txt":  _load_text,
    ".pdf":  _load_pdf,
    ".docx": _load_docx,
    ".xlsx": _load_xlsx,
}

# Extensions that are always active (xlsx only active when flag set)
SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}
if ENABLE_XLSX:
    SUPPORTED_EXTENSIONS.add(".xlsx")


def load_document(path: str) -> dict | None:
    """
    Load a document at the given path using the appropriate loader.
    Returns {text, metadata} or None if unsupported / unreadable.
    """
    ext = os.path.splitext(path)[1].lower()
    loader = _LOADERS.get(ext)
    if loader is None:
        logger.debug("no loader for extension %s (%s)", ext, path)
        return None
    return loader(path)
